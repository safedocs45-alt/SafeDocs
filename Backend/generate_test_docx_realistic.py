# generate_test_docx_realistic.py
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('SafeDocs Realistic Demo - DOCX (INERT STUBS)', level=1)

doc.add_paragraph(
    "This is a harmless, inert demo file. It includes TEXT stubs that mimic real macro payloads. "
    "These do NOT execute. Marker: X-SAFEDOCS-MARKER: MAL_TEST_VBA_STUB"
)

# Add an inert pseudo-VBA block as text (not an actual macro)
p = doc.add_paragraph()
run = p.add_run("\n--- BEGIN PSEUDO-VBA-STUB ---\n")
run.bold = True
run.font.size = Pt(10)
p.add_run(
    "Sub AutoOpen()\n"
    "' PSEUDO-VBA: inert stub example for demo only (DO NOT RUN)\n"
    "' The following lines are plain text and harmless\n"
    "    'MsgBox \"This is a safe inert stub\"\n"
    "End Sub\n"
)
p.add_run("\n--- END PSEUDO-VBA-STUB ---\n")

# Add an embedded-object placeholder (as plain text)
doc.add_paragraph("Embedded object placeholder: OLE_PAYLOAD_STUB (inert)")

# Add a fake base64 blob labelled clearly as inert
doc.add_paragraph("BASE64_STUB: VGhpcyBpcyBhIGR1bW15IGJsb2Igc3R1YiA9IGluaGVyZW50Cg==")

# Custom property comment with marker
doc.core_properties.comments = "X-SAFEDOCS-MARKER: MAL_TEST_VBA_STUB"

fn = 'safedocs_realistic_docx.docx'
doc.save(fn)
print("Created", fn)
